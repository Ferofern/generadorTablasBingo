import random as rd
from docx.shared import Pt
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import tkinter as tk
from tkinter import messagebox
import time as tm

def obtener_entero_positivo_input(mensaje):
    valor = entrada.get()
    try:
        valor = int(valor)
        if valor <= 0:
            messagebox.showerror("Error", "Ingresa un valor entero positivo.")
        else:
            return valor
    except ValueError:
        messagebox.showerror("Error", "Ingresa un valor entero válido.")


def generadorUnico(a):
    current_time = tm.localtime()
    formatted_time = tm.strftime("%d%m%Y", current_time)
    listaLetras = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L',
                   'M', 'N', 'O', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X',
                   'Y', 'Z']

    conjuntoCodigos = set()
    while len(conjuntoCodigos) < a:
        codigo = (
                         str(rd.randint(0, 9)) + rd.choice(listaLetras) +
                         str(rd.randint(0, 9)) + rd.choice(listaLetras) +
                         str(rd.randint(0, 9)) + rd.choice(listaLetras) +
                         str(rd.randint(0, 9)) + rd.choice(listaLetras)
                 ) + '-' + formatted_time
        conjuntoCodigos.add(codigo)
    listaCodigos = list(conjuntoCodigos)
    return listaCodigos

def set_vertical_alignment(cell, alignment):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), alignment)
    tcPr.append(vAlign)

def generadorTablas(document, archivo_csv, codigo):
    usados = set()
    listas = []
    for _ in range(5):
        lista = []
        for rango in [(1, 20), (21, 40), (41, 60), (61, 80), (81, 99)]:
            while True:
                numero = str(rd.randint(*rango))
                if numero not in usados:
                    usados.add(numero)
                    lista.append(numero)
                    break
        listas.append(lista)
    linea_csv = ';'.join([codigo] + [','.join(l) for l in listas])

    with open(archivo_csv, mode='a', newline='') as file:
        file.write(linea_csv + '\n')

    bingo = ['B', 'I', 'N', 'G', 'O']
    mensaje = 'Tabla: ' + codigo
    document.add_heading(mensaje, level=1)

    # Insertamos un párrafo al principio del documento para posicionar la tabla
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break()

    # Añadimos la tabla después del párrafo creado
    tablelistaAOC = document.add_table(rows=6, cols=len(bingo), style='Light Grid Accent 4')

    widths = [0.5, 0.5, 0.5, 0.5, 0.5]
    for row in tablelistaAOC.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Pt(width * 72)
    heights = [0.5, 0.5, 0.5, 0.5, 0.5, 0.5]
    for row, height in zip(tablelistaAOC.rows, heights):
        row.height = Pt(height * 72)

    for i in range(len(bingo)):
        cell = tablelistaAOC.cell(0, i)
        cell.text = bingo[i]
        run = cell.paragraphs[0].runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(15)
        run.bold = True
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        set_vertical_alignment(cell, 'center')

    for idx, lista in enumerate(listas):
        for j in range(len(lista)):
            cell = tablelistaAOC.cell(idx + 1, j)
            cell.text = lista[j]
            run = cell.paragraphs[0].runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(15)
            run.bold = False
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            set_vertical_alignment(cell, 'center')

def generar_tablas():
    cuantasTablas = obtener_entero_positivo_input('Cuantas tablas desea crear: ')
    document = Document()
    archivo_csv = rutaCSVDeTablas

    # Obtener todos los códigos únicos necesarios de una vez
    codigos_unicos = generadorUnico(cuantasTablas)

    # Iterar para crear las tablas
    for codigo in codigos_unicos:
        generadorTablas(document, archivo_csv, codigo)

    ruta_completa = rutaCSVDeTablasGanadoras
    document.save(ruta_completa)
    messagebox.showinfo("Terminado", "Proceso completado. Se ha guardado el archivo en: BingoStarDist/TablasImprimir.docx")


# rutas

rutaCSVDeTablas = 'C:/Users/ferof/OneDrive/Desktop/BingoStarDist/BufferBorrarSoloElContenido/Tablas.csv'
rutaCSVDeTablasGanadoras = 'C:/Users/ferof/OneDrive/Desktop/BingoStarDist/BufferBorrarSoloElContenido/TablasImprimir.docx'
rutaIco = 'C:/Users/ferof/OneDrive/Desktop/BingoStarDist/Imagenes/CompanyLogo.ico'
rutaPenege = 'C:/Users/ferof/OneDrive/Desktop/BingoStarDist/Imagenes/CompanyLogo.png'
# Crear la ventana principal
root = tk.Tk()
root.title("BingoStar")
root.title('BingoStar 1.0')
root.geometry('400x150')
root.iconbitmap(rutaIco)

# Crear etiqueta y entrada para el número de tablas
tk.Label(root, text="Número de tablas:").pack()
entrada = tk.Entry(root)
entrada.pack()

# Crear botón para generar las tablas
tk.Button(root, text="Generar Tablas", command=generar_tablas).pack()
# Configuración adicional de la interfaz
Company = tk.Label(root, text='Software Desarrollado por: RedStar Developers ✫', font=('Times', 10))
Company.place(x=10, y=70)

CompanyPhoto = tk.PhotoImage(file=rutaPenege)
CompanyPhotoPantalla = tk.Label(root, image=CompanyPhoto)
CompanyPhotoPantalla.place(x=299, y=10)

root.mainloop()
